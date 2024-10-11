import PyPDF2
from docx import Document
from docx.shared import Inches

# Function to extract text from a PDF
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() or ""  # Ensure we avoid None type if no text is extracted
    return text

# Function to format the extracted text into sections
def format_resume_text(raw_text):
    lines = raw_text.split('\n')
    
    # Initialize structured output
    resume_data = {
        "SKILLS": "",
        "INTERESTS": "",
        "EDUCATION": "",
        "INDUSTRIAL PROJECT": "",
        "ACHIEVEMENTS": "",
        "ORGANIZATION": "",
        "CERTIFICATES": "",
        "PROJECTS": ""
    }
    
    current_section = None

    for line in lines:
        if line.strip() == "":
            continue

        line_upper = line.upper()  # Convert line to uppercase for comparison

        # Check for section headers
        for section in resume_data.keys():
            if section in line_upper:  # Check for section headers without braces
                current_section = section
                break

        # Append the line to the appropriate section content
        if current_section:
            resume_data[current_section] += line.strip() + "\n"
    
    return resume_data

# Function to insert a table after a specific paragraph without adding headers again
def insert_table_after_paragraph(doc, para, left_content, right_content):
    table = doc.add_table(rows=1, cols=2)  # Create the table after the paragraph
    table.autofit = False
    table.columns[0].width = Inches(3)  # Left side (half page)
    table.columns[1].width = Inches(3)  # Right side (half page)

    # Add left content (no header)
    table.cell(0, 0).text = left_content

    # Add right content (no header)
    table.cell(0, 1).text = right_content

    # Move the table after the paragraph
    para._element.addnext(table._element)

# Function to format points into lines (two points per line)
def format_points(content):
    points = content.strip().split('\n')  # Split by new lines to get individual points
    formatted_lines = []

    # Group points into pairs for formatting
    for i in range(0, len(points), 2):
        line = ""
        if i < len(points):
            line += points[i].strip()  # First point
        if i + 1 < len(points):
            line += " | " + points[i + 1].strip()  # Second point (if exists)
        formatted_lines.append(line)

    return '\n'.join(formatted_lines)

# Function to fill the Word template dynamically
def fill_docx_template(template_path, output_path, resume_data):
    doc = Document(template_path)

    # Debug: Print all paragraphs in the template
    for para in doc.paragraphs:
        print(f"Paragraph: {para.text}")  # Print each paragraph for debugging

    # Traverse the paragraphs in the document
    for para in doc.paragraphs:
        found_headers = [key for key in resume_data.keys() if key in para.text]
        
        # Debugging: Print detected headers
        print(f"Found headers in paragraph: {found_headers}")

        # Handle case where two headers are on the same line
        if len(found_headers) == 2:  # If two headers are found on the same line
            left_header, right_header = found_headers
            left_content = resume_data[left_header].strip()
            right_content = resume_data[right_header].strip()

            # Clear the placeholders, but do not repeat the headers
            para.clear()  # Clear the paragraph content but keep the headers intact
            insert_table_after_paragraph(doc, para, left_content, right_content)  # Insert content

        # For single header, replace the placeholder with content
        for key in resume_data.keys():
            placeholder = key  # For example, SKILLS
            if placeholder in para.text:
                # Format points if the content has multiple lines
                formatted_content = format_points(resume_data[key].strip())
                print(f"Replacing {placeholder} with: {formatted_content}")  # Debug output
                para.text = para.text.replace(placeholder, formatted_content)

    # Save the updated document
    doc.save(output_path)

# Function to check and validate the structure of the template
def validate_template(template_path, resume_data):
    doc = Document(template_path)

    missing_placeholders = []
    # Ensure all required placeholders are present in the document
    for key in resume_data.keys():
        found = False
        for para in doc.paragraphs:
            placeholder = key  # For example, SKILLS
            if placeholder in para.text:
                found = True
                break
        if not found:
            missing_placeholders.append(key)

    if missing_placeholders:
        print(f"Missing placeholders: {missing_placeholders}")
    else:
        print("All placeholders are correctly placed in the template.")

# Main part: extracting and populating resume
pdf_path = 'Final.pdf'  # Replace with your PDF file path
template_path = 'My Resume (6).docx'  # Path to your Word template file
output_path = 'populated_resume.docx'  # Define the output Word file path

# Extract the text from the PDF resume
resume_text = extract_text_from_pdf(pdf_path)

# Format the extracted resume text
formatted_resume = format_resume_text(resume_text)

# Check and validate the template to ensure placeholders are correct
validate_template(template_path, formatted_resume)

# Fill the Word template with the formatted resume data
fill_docx_template(template_path, output_path, formatted_resume)

print(f"\nPopulated resume has been saved to {output_path}.")
